# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "TradeTest_nav_cloud_next_steps.docx"


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_east_asia_font(r)
    r.font.size = Pt(9.5)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        shade_cell(cell, "0B2545")
        cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_east_asia_font(run)
        run.font.color.rgb = RGBColor(11, 37, 69)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_east_asia_font(r1)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        set_east_asia_font(r2)
    else:
        r = p.add_run(text)
        set_east_asia_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_east_asia_font(r)
    return p


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TradeTest：左侧导航、云端保存与下一步功能方案")
    set_east_asia_font(run)
    run.bold = True
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("供产品决策使用：本文件只覆盖第 1、2、5 点，标的搜索确认与组合输入已进入实现。")
    set_east_asia_font(r)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 116, 139)

    add_heading(doc, "一、左侧导航：哪些保留、隐藏、以及怎么用", 1)
    add_para(doc, "推荐原则：左侧导航只放用户能立即理解并且有真实功能承接的入口。没有功能的入口先隐藏，不要制造“点了也没用”的落差。")
    add_table(
        doc,
        ["模块", "建议", "用途", "第一版处理"],
        [
            ["总览", "保留", "显示最近回测、累计运行次数、最近收益/回撤摘要。", "做成 Dashboard 首页。"],
            ["新建回测", "保留", "当前主工作台，创建单标的或组合回测。", "继续作为默认入口。"],
            ["回测列表", "保留", "展示云端和本地保存的历史记录。", "接云端 backtest_runs。"],
            ["策略库", "保留", "展示可用策略和下一版策略。", "MA/RSI 可套用，Breakout/Bollinger 标注下一版。"],
            ["数据管理", "保留", "股票搜索、已确认标的、数据源状态、数据质量提示。", "承接搜索确认逻辑。"],
            ["回测报告", "保留", "集中下载 Excel、查看图表、重新打开 YAML。", "从历史记录拆出下载入口。"],
            ["设置", "保留", "云端配置、默认市场、保存上限、语言设置。", "云端大表单移到这里。"],
            ["收益分析", "暂时隐藏", "需要多次回测对比后才有意义。", "后续有对比图再开放。"],
            ["因子分析", "暂时隐藏", "当前没有因子数据和因子模型。", "避免空入口。"],
            ["组合分析", "改名/后置", "建议改成“组合回测”，放到回测组。", "组合模式稳定后再显示。"],
            ["收藏策略", "暂时隐藏", "策略数量少时价值不高。", "策略库成熟后再开放。"],
        ],
        [1.05, 0.9, 3.2, 2.0],
    )

    add_heading(doc, "建议的新排布", 2)
    for item in [
        "总览",
        "回测：新建回测 / 回测列表 / 回测报告",
        "策略与数据：策略库 / 数据管理 / 组合回测",
        "系统：设置",
        "最底部小卡片：云端同步状态、终端 ID、配置名称。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "二、云端保存：压到底部，并控制保存上限", 1)
    add_para(doc, "推荐原则：云端是基础设施，不应该占据主操作区。它只需要告诉用户“是否已连接”，详细配置放到设置页。")
    add_table(
        doc,
        ["保存内容", "建议保存数量", "说明"],
        [
            ["YAML + summary 指标", "最近 1000 条", "体积小，便于回溯每次运行配置。"],
            ["Excel 报告 + PNG 图表", "最近 300 条", "附件体积更大，超过后保留轻量记录即可。"],
            ["超过上限的旧记录", "自动清理", "避免免费额度被不知不觉打满。"],
        ],
        [2.0, 1.4, 3.6],
    )
    add_para(doc, "实现建议：每次点击“运行回测”时自动暂存 YAML。运行成功后再保存 summary、Excel、PNG；运行失败则只保存失败 YAML 和错误摘要，方便复盘。")
    add_para(doc, "容量判断：Supabase 免费层主要受数据库和文件存储限制影响。1000 条 YAML/summary 通常很安全；真正占空间的是 Excel 和 PNG，因此附件建议设置 300 条软上限。")

    add_heading(doc, "五、下一步功能顺序建议", 1)
    add_para(doc, "推荐先把“输入和确认”做稳，再扩展左侧导航和分析页。这样用户不会因为输错标的、市场混用或组合权重错误得到误导性结果。")
    add_table(
        doc,
        ["优先级", "功能", "为什么先/后做", "验收标准"],
        [
            ["P0", "标的搜索与确认", "避免用户输入 qqq 就直接跑错标的。", "必须确认候选后才能运行。"],
            ["P0", "单标的 / 组合模式", "这是回测逻辑分岔点，必须先定。", "单标的走 MA/RSI，组合走权重调仓。"],
            ["P1", "云端轻量保存策略", "每次运行都留下 YAML，可追溯。", "最近 1000 条 YAML 可查看。"],
            ["P1", "左侧导航真实化", "减少无效按钮，提升可信度。", "所有可见入口都有明确页面或说明。"],
            ["P2", "组合分析页", "需要历史数据积累后才有价值。", "支持组合收益、回撤、持仓和调仓明细。"],
            ["P2", "策略收藏", "策略库丰富后再做。", "能收藏模板并一键套用。"],
        ],
        [0.8, 1.7, 2.5, 2.2],
    )

    add_heading(doc, "待你确认的问题", 1)
    for item in [
        "左侧是否保留“收益分析 / 因子分析 / 收藏策略”，还是先隐藏？",
        "云端是否采用“1000 条轻量记录 + 300 条完整附件”的默认上限？",
        "组合回测入口是放在“新建回测”的标的模式里，还是单独放到左侧“组合回测”？",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
